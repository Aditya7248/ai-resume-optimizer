"""
ATS Score Computation — 6-signal weighted model.

Grounded in real ATS behavior research:
- Workday Skills Cloud (NLP-based, 200k+ canonical skills)
- Taleo (strict literal matching, title weighting)
- Greenhouse (section completeness, human-readable)
- Lever (semantic matching, prose quality)

Score breakdown (max 100):
  keyword_match        — 30 pts
  section_completeness — 20 pts
  format_parsability   — 20 pts
  keyword_placement    — 15 pts
  date_consistency     — 10 pts
  file_health          —  5 pts
"""

import io
import re
from models.resume import ParsedResume
from models.jd import ParsedJD
from models.optimization import ATSScoreBreakdown


# ── Signal 1: Keyword Match (30 pts) ─────────────────────────────────────────

import re as _re_ats
import functools

_STRIP_RE = _re_ats.compile(r"[^a-z0-9\s]")


@functools.lru_cache(maxsize=1024)
def _word_variants(word: str) -> frozenset[str]:  # frozenset is hashable → enables lru_cache
    """
    Generate common morphological variants of a word so we can match
    "pipeline" ↔ "pipelines", "engineer" ↔ "engineered" ↔ "engineering",
    "deploy" ↔ "deployment" ↔ "deployed", "manage" ↔ "management" ↔ "manager", etc.
    Works for any domain — tech, HR, finance, management, healthcare.
    """
    variants = {word}

    def _add_forms(base: str):
        """Add all inflectional forms for a given base word."""
        variants.add(base)
        variants.add(base + "s")        # deploy → deploys
        variants.add(base + "es")       # process → processes
        variants.add(base + "d")        # code → coded
        variants.add(base + "ed")       # deploy → deployed
        variants.add(base + "ing")      # deploy → deploying
        variants.add(base + "r")        # manage → manager
        variants.add(base + "er")       # build → builder
        variants.add(base + "ment")     # manage → management
        variants.add(base + "tion")     # implement → implementation (loose)
        variants.add(base + "ation")    # transform → transformation
        variants.add(base + "al")       # profession → professional
        variants.add(base + "ly")       # professional → professionally

    # Add forms for the original word
    _add_forms(word)

    # Strip common suffixes → derive root → add forms for root too
    for suffix in ("ation", "tion", "ment", "ing", "ed", "er", "al", "ly", "s"):
        if word.endswith(suffix) and len(word) - len(suffix) >= 3:
            root = word[:-len(suffix)]
            _add_forms(root)
            _add_forms(root + "e")   # manag → manage → managed, manager, management…

    return frozenset(variants)


def _keyword_in_text(keyword: str, resume_text: str, resume_words: set[str]) -> bool:
    """
    Flexible keyword matching — handles plurals, word-order variation, and common suffixes.
    Works for any domain: tech, HR, finance, management, etc.
    Strategy:
      1. Literal substring (fastest, exact match)
      2. Multi-word: every keyword word has at least one variant present in resume
         (handles "data pipeline" ↔ "data pipelines", any word order)
      3. Single-word: any variant of keyword found in resume word set
    """
    # 1. Literal match (fastest)
    if keyword in resume_text:
        return True

    # Normalise keyword (remove punctuation, split into words)
    kw_clean = _STRIP_RE.sub(" ", keyword).strip()
    kw_words = kw_clean.split()

    if not kw_words:
        return False

    # 2. Multi-word: every keyword word (by variants) is present in resume word set
    if len(kw_words) > 1:
        if all(
            bool(_word_variants(w) & resume_words)  # at least one variant in resume
            for w in kw_words if len(w) >= 3
        ):
            return True

    # 3. Single-word: any variant of keyword in resume word set
    if len(kw_words) == 1 and len(kw_words[0]) >= 3:
        if _word_variants(kw_words[0]) & resume_words:
            return True

    return False


def _compute_keyword_match(resume: ParsedResume, jd: ParsedJD) -> float:
    resume_text = (
        " ".join(resume.skills) + " " +
        " ".join(resume.tech_stack) + " " +
        " ".join([b for exp in resume.experience for b in exp.bullets]) + " " +
        (resume.summary or "")
    ).lower()

    # Pre-compute word set for fast containment checks
    import re as _re
    resume_words = set(_re.sub(r"[^a-z0-9\s]", " ", resume_text).split())

    hard_kw = [k.lower() for k in jd.hard_keywords]
    soft_kw = [k.lower() for k in jd.soft_keywords]

    if not hard_kw and not soft_kw:
        return 15.0  # no keywords extracted → neutral score

    # Hard keywords weighted 2x
    hard_matched = sum(1 for k in hard_kw if _keyword_in_text(k, resume_text, resume_words))
    soft_matched = sum(1 for k in soft_kw if _keyword_in_text(k, resume_text, resume_words))

    hard_score = (hard_matched / max(len(hard_kw), 1)) * 20
    soft_score = (soft_matched / max(len(soft_kw), 1)) * 10

    return min(round(hard_score + soft_score, 1), 30.0)


# ── Signal 2: Section Completeness (20 pts) ──────────────────────────────────

def _compute_section_completeness(resume: ParsedResume) -> float:
    score = 0.0

    # Contact info (4 pts)
    pi = resume.personal_info
    if pi.full_name and pi.full_name != "Unknown":
        score += 1
    if pi.email:
        score += 1.5
    if pi.phone:
        score += 1.5

    # Summary (4 pts)
    if resume.summary and len(resume.summary) > 30:
        score += 4

    # Experience with dates (4 pts)
    if resume.experience:
        has_dates = all(
            exp.start_date and exp.end_date
            for exp in resume.experience
        )
        score += 4 if has_dates else 2

    # Skills section (4 pts)
    if len(resume.skills) >= 3:
        score += 4
    elif resume.skills:
        score += 2

    # Education with degree + year (4 pts)
    if resume.education:
        has_year = all(edu.year for edu in resume.education)
        score += 4 if has_year else 2

    return min(round(score, 1), 20.0)


# ── Signal 3: Format Parsability (20 pts) ────────────────────────────────────
# Evaluated on raw bytes for DOCX; for PDF we use heuristics

def _compute_format_parsability(file_bytes: bytes, content_type: str) -> float:
    score = 20.0  # start full, deduct for issues

    if "pdf" in content_type.lower():
        # PDFs are harder to assess structurally without rendering
        # Deduct if text extraction was very short (possible image-based)
        if len(file_bytes) > 500_000:  # large file with presumably scanned content
            score -= 5
        return max(score, 10.0)

    # DOCX inspection
    try:
        from docx import Document
        import io as _io
        doc = Document(_io.BytesIO(file_bytes))

        # Penalize tables used for layout (common in complex CV templates)
        if len(doc.tables) > 3:
            score -= 4

        # Penalize text boxes / drawing objects
        body_xml = doc.element.body.xml
        if "txbx" in body_xml or "textbox" in body_xml.lower():
            score -= 4

        # Penalize headers/footers with critical info
        for section in doc.sections:
            header_text = " ".join(p.text for p in section.header.paragraphs)
            if any(word in header_text.lower() for word in ["name", "email", "phone"]):
                score -= 2

        # Penalize two-column layouts (common Workday parsing failure)
        # Detected via section column count
        for section in doc.sections:
            if section._sectPr is not None:
                cols = section._sectPr.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols"
                )
                if cols:
                    num = cols[0].get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
                    )
                    if num and int(num) > 1:
                        score -= 5
                        break

    except Exception:
        score = 12.0  # couldn't inspect, give partial score

    return max(round(score, 1), 0.0)


# ── Signal 4: Keyword Placement (15 pts) ─────────────────────────────────────

def _compute_keyword_placement(resume: ParsedResume, jd: ParsedJD) -> float:
    import re as _re
    hard_kw = [k.lower() for k in jd.hard_keywords[:15]]
    if not hard_kw:
        return 8.0  # neutral

    # Keywords in summary carry highest weight
    summary_text = (resume.summary or "").lower()
    summary_words = set(_re.sub(r"[^a-z0-9\s]", " ", summary_text).split())
    summary_hits = sum(1 for k in hard_kw if _keyword_in_text(k, summary_text, summary_words))

    # Skills section keywords
    skills_text = " ".join(resume.skills + resume.tech_stack).lower()
    skills_words = set(_re.sub(r"[^a-z0-9\s]", " ", skills_text).split())
    skills_hits = sum(1 for k in hard_kw if _keyword_in_text(k, skills_text, skills_words))

    summary_score = (summary_hits / max(len(hard_kw), 1)) * 8
    skills_score = (skills_hits / max(len(hard_kw), 1)) * 7

    return min(round(summary_score + skills_score, 1), 15.0)


# ── Signal 5: Date Consistency (10 pts) ──────────────────────────────────────

def _compute_date_consistency(resume: ParsedResume) -> float:
    if not resume.experience:
        return 5.0  # no experience to check

    score = 10.0

    date_formats_seen = set()
    # Detect date format variations
    month_year_pattern = re.compile(
        r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}",
        re.IGNORECASE
    )
    mm_yyyy_pattern = re.compile(r"\d{2}/\d{4}")
    yyyy_pattern = re.compile(r"^\d{4}$")

    for exp in resume.experience:
        for date_str in [exp.start_date, exp.end_date]:
            if not date_str or date_str.lower() == "present":
                continue
            if month_year_pattern.search(date_str):
                date_formats_seen.add("month_year")
            elif mm_yyyy_pattern.search(date_str):
                date_formats_seen.add("mm_yyyy")
            elif yyyy_pattern.match(date_str.strip()):
                date_formats_seen.add("yyyy")

    # Multiple date formats used = deduct
    if len(date_formats_seen) > 1:
        score -= 4

    # Check chronological order (most recent first)
    years = []
    for exp in resume.experience:
        year_match = re.search(r"\d{4}", exp.start_date or "")
        if year_match:
            years.append(int(year_match.group()))

    if years and years != sorted(years, reverse=True):
        score -= 3  # not reverse chronological

    return max(round(score, 1), 0.0)


# ── Signal 6: File Health (5 pts) ────────────────────────────────────────────

def _compute_file_health(file_bytes: bytes, content_type: str, parsed_resume: ParsedResume) -> float:
    score = 5.0

    # Text was extractable (not image-based)
    if not parsed_resume.raw_text or len(parsed_resume.raw_text) < 100:
        return 0.0  # total failure

    # Check for garbled unicode
    raw = parsed_resume.raw_text or ""
    garbled = sum(1 for c in raw if ord(c) > 65000)
    if garbled > 10:
        score -= 2

    return max(round(score, 1), 0.0)


# ── Master Scorer ─────────────────────────────────────────────────────────────

def compute_ats_score(
    resume_bytes: bytes,
    content_type: str,
    parsed_resume: ParsedResume,
    parsed_jd: ParsedJD,
) -> ATSScoreBreakdown:
    keyword_match = _compute_keyword_match(parsed_resume, parsed_jd)
    section_completeness = _compute_section_completeness(parsed_resume)
    format_parsability = _compute_format_parsability(resume_bytes, content_type)
    keyword_placement = _compute_keyword_placement(parsed_resume, parsed_jd)
    date_consistency = _compute_date_consistency(parsed_resume)
    file_health = _compute_file_health(resume_bytes, content_type, parsed_resume)

    total = (
        keyword_match + section_completeness + format_parsability +
        keyword_placement + date_consistency + file_health
    )

    return ATSScoreBreakdown(
        keyword_match=keyword_match,
        section_completeness=section_completeness,
        format_parsability=format_parsability,
        keyword_placement=keyword_placement,
        date_consistency=date_consistency,
        file_health=file_health,
        total=round(total, 1),
    )


def compute_ats_score_from_text(
    rewritten_resume: dict,
    jd: ParsedJD,
    original_resume: "ParsedResume | None" = None,
) -> "ATSScoreBreakdown":
    """
    Post-optimization ATS score computed from rewritten content dict.
    Used after AI rewriting to show improvement.

    Pass `original_resume` to carry over personal_info, education, and dates
    that are never rewritten — this prevents the "after" score from losing
    points for contact info that was trimmed from the slim rewrite prompt.
    """
    from models.resume import ParsedResume, PersonalInfo

    # Contact info: prefer original (never rewritten) to avoid slim-prompt gaps
    if original_resume is not None:
        personal_info = original_resume.personal_info
    else:
        pi_data = rewritten_resume.get("personal_info", {})
        personal_info = PersonalInfo(
            full_name=pi_data.get("full_name", ""),
            email=pi_data.get("email"),
            phone=pi_data.get("phone"),
        )

    from models.resume import ExperienceEntry, EducationEntry
    experience = []
    for exp in rewritten_resume.get("experience", []):
        try:
            experience.append(ExperienceEntry(**exp))
        except Exception:
            pass

    # Education: prefer original (never rewritten); keeps GPA/year for section scoring
    if original_resume is not None and original_resume.education:
        education = original_resume.education
    else:
        education = []
        for edu in rewritten_resume.get("education", []):
            try:
                education.append(EducationEntry(**edu))
            except Exception:
                pass

    resume = ParsedResume(
        personal_info=personal_info,
        summary=rewritten_resume.get("summary"),
        skills=rewritten_resume.get("skills", []),
        experience=experience,
        education=education,
        tech_stack=rewritten_resume.get("tech_stack", []),
        raw_text=" ".join([
            rewritten_resume.get("summary", ""),
            " ".join(rewritten_resume.get("skills", [])),
            " ".join([b for exp in experience for b in exp.bullets]),
        ]),
    )

    # Format parsability is 20 for rewritten content (clean text output)
    keyword_match = _compute_keyword_match(resume, jd)
    section_completeness = _compute_section_completeness(resume)
    keyword_placement = _compute_keyword_placement(resume, jd)
    date_consistency = _compute_date_consistency(resume)

    total = keyword_match + section_completeness + 20.0 + keyword_placement + date_consistency + 5.0

    return ATSScoreBreakdown(
        keyword_match=keyword_match,
        section_completeness=section_completeness,
        format_parsability=20.0,
        keyword_placement=keyword_placement,
        date_consistency=date_consistency,
        file_health=5.0,
        total=round(total, 1),
    )
