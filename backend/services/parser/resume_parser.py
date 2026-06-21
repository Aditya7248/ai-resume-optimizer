import io
import re
from typing import Optional

import pdfplumber
from docx import Document

from models.resume import (
    ParsedResume, PersonalInfo, ExperienceEntry,
    EducationEntry, ProjectEntry, CertificationEntry,
)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def detect_language(text: str) -> str:
    """
    Detect language using langdetect library (supports 55+ languages).
    Falls back to keyword heuristic if langdetect fails or is unavailable.
    """
    try:
        from langdetect import detect, DetectorFactory
        DetectorFactory.seed = 0  # deterministic results
        lang = detect(text[:3000])  # first 3000 chars is enough
        return lang
    except Exception:
        # Fallback: keyword heuristic for English
        english_markers = ["experience", "education", "skills", "summary", "work", "project"]
        text_lower = text.lower()
        hits = sum(1 for word in english_markers if word in text_lower)
        return "en" if hits >= 2 else "unknown"


def extract_email(text: str) -> Optional[str]:
    match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    match = re.search(r"(\+?\d[\d\s\-().]{8,15}\d)", text)
    return match.group(0).strip() if match else None


def extract_linkedin(text: str) -> Optional[str]:
    match = re.search(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    return f"https://{match.group(0)}" if match else None


def extract_github(text: str) -> Optional[str]:
    match = re.search(r"github\.com/[\w\-]+", text, re.IGNORECASE)
    return f"https://{match.group(0)}" if match else None


def estimate_total_experience(text: str) -> Optional[float]:
    """
    Looks for patterns like 'X+ years', 'X years of experience'
    and also calculates from date ranges if present.
    """
    # Direct mention
    pattern = re.search(
        r"(\d+(?:\.\d+)?)\s*\+?\s*years?\s*(?:of\s+)?(?:experience|exp)",
        text, re.IGNORECASE
    )
    if pattern:
        return float(pattern.group(1))

    # Try to find date ranges (e.g., Jan 2020 - Present)
    date_pattern = re.findall(
        r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|"
        r"April|June|July|August|September|October|November|December)\s+(\d{4})",
        text, re.IGNORECASE
    )
    if len(date_pattern) >= 2:
        years = sorted([int(d[1]) for d in date_pattern])
        if years:
            return float(years[-1] - years[0])
    return None


def parse_resume(file_bytes: bytes, content_type: str, filename: str) -> ParsedResume:
    """
    Extracts raw text from PDF or DOCX and builds a ParsedResume.
    AI-based deep extraction happens in ai_service.py — this handles
    the deterministic extraction layer.
    """
    # Extract raw text
    if "pdf" in content_type.lower():
        raw_text = extract_text_from_pdf(file_bytes)
    else:
        raw_text = extract_text_from_docx(file_bytes)

    if not raw_text or len(raw_text.strip()) < 50:
        raise ValueError(
            "Could not extract text from the resume. "
            "The file may be scanned/image-based or corrupted. "
            "Please upload a text-based PDF or DOCX."
        )

    # Language check
    lang = detect_language(raw_text)
    if lang != "en":
        raise ValueError(
            "Resume language could not be identified as English. "
            "This tool currently supports English-language resumes only."
        )

    # Extract personal info deterministically
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    linkedin = extract_linkedin(raw_text)
    github = extract_github(raw_text)

    # Extract first non-empty line as name heuristic
    lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
    name = lines[0] if lines else "Unknown"

    # Estimate total experience
    total_exp = estimate_total_experience(raw_text)

    personal_info = PersonalInfo(
        full_name=name,
        email=email,
        phone=phone,
        linkedin=linkedin,
        github=github,
    )

    # Return partial resume — AI service will enrich with structured data
    return ParsedResume(
        personal_info=personal_info,
        total_years_experience=total_exp,
        raw_text=raw_text,
    )
