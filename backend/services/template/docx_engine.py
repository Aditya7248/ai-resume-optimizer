"""
DOCX Template Engine — Run-level text injection.

Core principle: ONLY run.text is ever written.
All style properties (font, size, color, bold, italic, spacing) are NEVER touched.
This preserves the exact visual design of the user's uploaded template.
"""

import io
import os
import re
import subprocess
from copy import deepcopy
from typing import Any

from docx import Document
from docx.oxml.ns import qn

OUTPUT_DIR = "/tmp/resume-optimizer"


# ─── Keep-My-Format: direct text injection into user's own DOCX ──────────────

def _normalize_bullet(text: str) -> str:
    """Strip common bullet/list characters so we can match against parsed text."""
    return re.sub(r'^[•\-\*·→⟩▪▸◆◉]\s*', '', text).strip()


def _build_text_map(original_resume: dict, rewritten_resume: dict) -> dict[str, str]:
    """
    Build original_text → rewritten_text for every piece of content that changed.
    Only maps content that actually differs so unchanged paragraphs are untouched.
    """
    text_map: dict[str, str] = {}

    # Summary
    orig = (original_resume.get("summary") or "").strip()
    new = (rewritten_resume.get("summary") or "").strip()
    if orig and new and orig != new:
        text_map[orig] = new

    # Experience bullets
    for oe, ne in zip(
        original_resume.get("experience", []),
        rewritten_resume.get("experience", []),
    ):
        for ob, nb in zip(
            oe.get("bullets", []) if isinstance(oe, dict) else [],
            ne.get("bullets", []) if isinstance(ne, dict) else [],
        ):
            ob_c = ob.strip() if ob else ""
            nb_c = nb.strip() if nb else ""
            if ob_c and nb_c and ob_c != nb_c:
                text_map[ob_c] = nb_c
                # Also map the normalised form (in case DOCX has leading bullet char)
                ob_norm = _normalize_bullet(ob_c)
                if ob_norm and ob_norm != ob_c:
                    text_map[ob_norm] = nb_c

    # Project bullets
    for op, np_ in zip(
        original_resume.get("projects", []),
        rewritten_resume.get("projects", []),
    ):
        orig_desc = (op.get("description") or "").strip() if isinstance(op, dict) else ""
        new_desc = (np_.get("description") or "").strip() if isinstance(np_, dict) else ""
        if orig_desc and new_desc and orig_desc != new_desc:
            text_map[orig_desc] = new_desc
        for ob, nb in zip(
            op.get("bullets", []) if isinstance(op, dict) else [],
            np_.get("bullets", []) if isinstance(np_, dict) else [],
        ):
            ob_c = ob.strip() if ob else ""
            nb_c = nb.strip() if nb else ""
            if ob_c and nb_c and ob_c != nb_c:
                text_map[ob_c] = nb_c

    return text_map


def _replace_text_content(doc: Document, text_map: dict[str, str]):
    """
    Walk all paragraphs/cells. When a paragraph's text matches an original,
    rewrite only run.text values — all run formatting properties are untouched.
    """
    def try_para(para):
        if not para.runs:
            return
        full = "".join(r.text for r in para.runs)
        full_norm = _normalize_bullet(full.strip())

        for orig, new_val in text_map.items():
            if not orig:
                continue
            orig_norm = _normalize_bullet(orig)
            if full_norm == orig_norm or full.strip() == orig:
                # Preserve any leading bullet character
                prefix_m = re.match(r'^([•\-\*·→⟩▪▸◆◉]\s*)', full)
                prefix = prefix_m.group(1) if prefix_m else ""
                para.runs[0].text = prefix + new_val
                for run in para.runs[1:]:
                    run.text = ""
                return

    for para in doc.paragraphs:
        try_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    try_para(para)

    for section in doc.sections:
        for para in section.header.paragraphs:
            try_para(para)


def inject_into_resume_docx(
    resume_bytes: bytes,
    original_resume: dict,
    rewritten_resume: dict,
    session_id: str,
) -> dict[str, str]:
    """
    Inject rewritten content directly into the candidate's own DOCX file.
    ONLY paragraph text is changed — every font, color, spacing, and style
    property on every run is completely preserved.
    """
    out_dir = os.path.join(OUTPUT_DIR, session_id)
    os.makedirs(out_dir, exist_ok=True)

    doc = Document(io.BytesIO(resume_bytes))
    text_map = _build_text_map(original_resume, rewritten_resume)
    _replace_text_content(doc, text_map)

    docx_filename = "optimized_resume.docx"
    docx_path = os.path.join(out_dir, docx_filename)
    doc.save(docx_path)

    pdf_filename = None
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", out_dir, docx_path],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0:
            pdf_filename = "optimized_resume.pdf"
        else:
            print(f"LibreOffice PDF warning: {result.stderr}")
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        print(f"PDF conversion unavailable: {e}")

    return {"docx": docx_filename, "pdf": pdf_filename}


def _replace_in_paragraph(para, placeholder: str, new_value: str):
    """
    Replace placeholder in a paragraph while preserving ALL run formatting.
    Handles the case where a placeholder is split across multiple runs.
    """
    # Check full text first
    full_text = "".join(run.text for run in para.runs)
    if placeholder not in full_text:
        return False

    # Single run case (most common)
    for run in para.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, new_value)
            return True

    # Multi-run case — placeholder split across runs
    # Strategy: find start and end run, merge into first, clear others
    combined = ""
    start_idx = None
    for i, run in enumerate(para.runs):
        combined += run.text
        if start_idx is None and placeholder[:len(combined)] == combined[-len(placeholder[:len(combined)]):]:
            # Partial match started
            if placeholder in combined:
                # Complete match
                if start_idx is None:
                    start_idx = i
                # Replace in first matched run
                first_run = para.runs[start_idx]
                first_run.text = combined.replace(placeholder, new_value)
                # Clear the other runs that were part of the placeholder
                for j in range(start_idx + 1, i + 1):
                    para.runs[j].text = ""
                return True
            start_idx = i
        elif start_idx is not None:
            if placeholder in combined:
                first_run = para.runs[start_idx]
                first_run.text = combined.replace(placeholder, new_value)
                for j in range(start_idx + 1, i + 1):
                    para.runs[j].text = ""
                return True

    # Fallback: replace in full and put in first run
    if para.runs:
        para.runs[0].text = full_text.replace(placeholder, new_value)
        for run in para.runs[1:]:
            run.text = ""
        return True

    return False


def _replace_in_cell(cell, placeholder: str, new_value: str):
    """Replace placeholder in all paragraphs of a table cell."""
    for para in cell.paragraphs:
        _replace_in_paragraph(para, placeholder, new_value)


def _replace_all(doc: Document, replacements: dict[str, str]):
    """Replace all placeholders throughout the document."""
    for placeholder, value in replacements.items():
        if not value:
            value = ""

        # Body paragraphs
        for para in doc.paragraphs:
            _replace_in_paragraph(para, placeholder, value)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _replace_in_cell(cell, placeholder, value)

        # Headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                _replace_in_paragraph(para, placeholder, value)
            for para in section.footer.paragraphs:
                _replace_in_paragraph(para, placeholder, value)


def _clone_paragraph(para, parent):
    """Deep copy a paragraph XML element and insert after original."""
    clone = deepcopy(para._element)
    para._element.addnext(clone)
    return clone


def _build_replacements(resume_data: dict) -> dict[str, str]:
    """Build placeholder → value map from rewritten resume data."""
    pi = resume_data.get("personal_info", {})
    experience = resume_data.get("experience", [])
    education = resume_data.get("education", [])
    projects = resume_data.get("projects", [])
    certifications = resume_data.get("certifications", [])
    skills = resume_data.get("skills", [])

    replacements = {
        "{{FULL_NAME}}": pi.get("full_name", ""),
        "{{NAME}}": pi.get("full_name", ""),
        "{{EMAIL}}": pi.get("email", ""),
        "{{PHONE}}": pi.get("phone", ""),
        "{{LOCATION}}": pi.get("location", ""),
        "{{LINKEDIN}}": pi.get("linkedin", ""),
        "{{GITHUB}}": pi.get("github", ""),
        "{{PORTFOLIO}}": pi.get("portfolio", ""),
        "{{SUMMARY}}": resume_data.get("summary", ""),
        "{{SKILLS_LIST}}": " · ".join(skills[:20]),
        "{{SKILLS_COMMA}}": ", ".join(skills[:20]),
    }

    # Experience entries
    for i, exp in enumerate(experience[:10], 1):
        bullets = exp.get("bullets", [])
        replacements.update({
            f"{{{{EXP_{i}_TITLE}}}}": exp.get("title", ""),
            f"{{{{EXP_{i}_COMPANY}}}}": exp.get("company", ""),
            f"{{{{EXP_{i}_LOCATION}}}}": exp.get("location", ""),
            f"{{{{EXP_{i}_DATES}}}}": f"{exp.get('start_date', '')} – {exp.get('end_date', '')}",
            f"{{{{EXP_{i}_START}}}}": exp.get("start_date", ""),
            f"{{{{EXP_{i}_END}}}}": exp.get("end_date", ""),
        })
        for j, bullet in enumerate(bullets[:6], 1):
            replacements[f"{{{{EXP_{i}_BULLET_{j}}}}}"] = bullet

    # Education entries
    for i, edu in enumerate(education[:3], 1):
        replacements.update({
            f"{{{{EDU_{i}_DEGREE}}}}": edu.get("degree", ""),
            f"{{{{EDU_{i}_INSTITUTION}}}}": edu.get("institution", ""),
            f"{{{{EDU_{i}_LOCATION}}}}": edu.get("location", ""),
            f"{{{{EDU_{i}_YEAR}}}}": edu.get("year", ""),
        })

    # Projects
    for i, proj in enumerate(projects[:5], 1):
        replacements.update({
            f"{{{{PROJ_{i}_NAME}}}}": proj.get("name", ""),
            f"{{{{PROJ_{i}_DESC}}}}": proj.get("description", ""),
            f"{{{{PROJ_{i}_TECH}}}}": ", ".join(proj.get("technologies", [])),
        })

    # Certifications
    for i, cert in enumerate(certifications[:5], 1):
        replacements[f"{{{{CERT_{i}}}}}"] = f"{cert.get('name', '')} — {cert.get('issuer', '')} {cert.get('year', '')}".strip(" —")

    return replacements


def inject_into_docx(
    template_bytes: bytes,
    resume_data: dict,
    session_id: str,
) -> dict[str, str]:
    """
    Main entry point.
    Injects rewritten resume content into the user's DOCX template.
    Returns paths to output DOCX and PDF files.
    """
    out_dir = os.path.join(OUTPUT_DIR, session_id)
    os.makedirs(out_dir, exist_ok=True)

    doc = Document(io.BytesIO(template_bytes))
    replacements = _build_replacements(resume_data)

    # Replace all placeholders
    _replace_all(doc, replacements)

    # Clean up any remaining unreplaced placeholders
    all_placeholders = re.compile(r"\{\{[A-Z_0-9]+\}\}")
    for para in doc.paragraphs:
        full = "".join(r.text for r in para.runs)
        if all_placeholders.search(full):
            cleaned = all_placeholders.sub("", full)
            if para.runs:
                para.runs[0].text = cleaned
                for run in para.runs[1:]:
                    run.text = ""

    # Save DOCX
    docx_filename = "optimized_resume.docx"
    docx_path = os.path.join(out_dir, docx_filename)
    doc.save(docx_path)

    # Convert to PDF via LibreOffice headless
    pdf_filename = None
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", out_dir, docx_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            pdf_filename = "optimized_resume.pdf"
        else:
            print(f"LibreOffice PDF conversion warning: {result.stderr}")
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        print(f"PDF conversion unavailable: {e}")

    return {
        "docx": docx_filename,
        "pdf": pdf_filename,
    }
